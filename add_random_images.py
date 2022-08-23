from docx import Document
from docx.shared import Inches
import random
import os
import PIL
from PIL import Image

# ------------------- CONSTANTS -------------------- #
DOC_NAME = "11th Day Together Wishes.docx"
IMAGE_WIDTH = Inches(6.5)
IMAGE_HEIGHT = Inches(4.88)
DIRECTORY_AB_PATH = r'C:\Users\Administrator\Downloads\pexels'
RATIO = 4 / 3
RATIO_MULTIPLICATIVE_INVERSE = 1 / RATIO
WIDTH_STANDARD = 1280
HEIGHT_STANDARD = 960


# -------- GET ALL IMAGE FILES FROM A DIRECTORY ---- #
def get_all_files():
    dir_path = DIRECTORY_AB_PATH

    result = []

    for file in os.listdir(dir_path):
        if file.endswith('.jpg'):
            result.append(file)
    with open("files_name.txt", mode="w") as files_name_file:
        for line in result:
            files_name_file.write(line)
            files_name_file.write("\n")
    return result


IMAGE_FILES = get_all_files()


# ---- CREATE DOCUMENT OBJECT AND GIVE ATTRIBUTES --- #
document = Document(DOC_NAME)
paragraphs = document.paragraphs


# ---- FORMAT IMAGE --------------------------------- #
def format_image(image_path, image_to_crop_name):
    # source: https://dev.to/flynestor/crop-a-picture-automatically-with-a-fixed-ratio-python-example-1j1a
    with Image.open(image_path) as image:
        width, height = image.size
        if width / height == RATIO:
            if width != WIDTH_STANDARD:
                image_crop = image.resize((WIDTH_STANDARD, HEIGHT_STANDARD))
            else:
                image_crop = image
        elif width / height > RATIO:
            offset = int(abs(width - RATIO * height) / 2)
            image_crop1 = image.crop((offset, 0, width - offset, height))
            image_crop = image_crop1.resize((WIDTH_STANDARD, HEIGHT_STANDARD))
        else:
            offset = int(abs(RATIO_MULTIPLICATIVE_INVERSE * width - height) / 2)
            image_crop1 = image.crop((0, offset, width, height - offset))
            image_crop = image_crop1.resize((WIDTH_STANDARD, HEIGHT_STANDARD))
        image_crop_path = DIRECTORY_AB_PATH + "\\" + "cropped\\" + "cropped_" + image_to_crop_name
        image_crop.save(image_crop_path, format="jpeg")
    return image_crop_path


# ----- SEARCH SIGN (**) AND REPLACE WITH IMAGE ----- #
for paragraph in paragraphs:
    if paragraph.text == "**":
        paragraph.text = ""
        image_run = paragraph.add_run()
        image_name = random.choice(IMAGE_FILES)
        path = DIRECTORY_AB_PATH + "\\" + image_name
        path_formatted = format_image(path, image_name)
        image_run.add_picture(path_formatted, width=IMAGE_WIDTH, height=IMAGE_HEIGHT)
        IMAGE_FILES.remove(image_name)

document.save(DOC_NAME)

# --------------------------- EXPLANATION ---------- #
# Add doc file
# read paragraphs
# iterate each paragraph
# find ** in paragraph.text
# add a run. assign it to a variable (lets call it image_run)
# use expression image_run.add_picture(FILENAME, width=, height=
